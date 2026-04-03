"""
Builds docs/FARJAR_Company_Profile.docx — layout aligned to FARJAR_Profile_Updated PDF (8 pages).

Edit CONTENT / paths below, then run:
  pip install python-docx
  python scripts/build_company_profile_docx.py

Images use files under assets/; swap paths if you replace photos. Missing files are skipped with a note.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# Brand (matches site)
ORANGE = RGBColor(255, 132, 2)
DARK = RGBColor(22, 26, 29)


def _root() -> Path:
    return Path(__file__).resolve().parent.parent


def _set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def add_cover(doc: Document) -> None:
    logo = _root() / "assets" / "images" / "logo" / "farjar-logo-stacked.png"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    if logo.is_file():
        p.add_run().add_picture(str(logo), width=Inches(2.2))
    else:
        p.add_run("[Logo: assets/images/logo/farjar-logo-stacked.png]").italic = True

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("FARJAR Construction & Development LLC")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = DARK
    t.paragraph_format.space_before = Pt(16)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("Company Profile")
    r2.font.size = Pt(13)
    r2.font.italic = True
    t2.paragraph_format.space_after = Pt(6)


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = DARK
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)


def add_subheading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = ORANGE
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15


def add_stat_grid(doc: Document) -> None:
    """Four KPI tiles like the PDF cover stats."""
    tbl = doc.add_table(rows=1, cols=4)
    tbl.autofit = False
    stats = [
        ("15+", "Cities Across Oman"),
        ("3+", "Years of Excellence"),
        ("100%", "Client Satisfaction"),
        ("42+", "Projects Completed"),
    ]
    for i, (num, label) in enumerate(stats):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        _set_cell_shading(cell, "FFF3E8")
        p1 = cell.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(num)
        r1.bold = True
        r1.font.size = Pt(20)
        r1.font.color.rgb = ORANGE
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(label)
        r2.font.size = Pt(9)
        r2.font.color.rgb = DARK
        cell.width = Inches(1.45)
    doc.add_paragraph()


def add_mission_vision_block(doc: Document) -> None:
    """Two columns: Mission | Vision (PDF style)."""
    tbl = doc.add_table(rows=1, cols=2)
    left, right = tbl.rows[0].cells[0], tbl.rows[0].cells[1]
    for c in (left, right):
        c.width = Inches(3.15)

    def fill_cell(cell, title_line: str, subtitle: str, body: str) -> None:
        cell.text = ""
        p0 = cell.add_paragraph()
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        t = p0.add_run(title_line)
        t.bold = True
        t.font.size = Pt(10)
        t.font.color.rgb = ORANGE
        p0.paragraph_format.space_after = Pt(2)
        p1 = cell.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        s = p1.add_run(subtitle)
        s.bold = True
        s.font.size = Pt(12)
        s.font.color.rgb = DARK
        p1.paragraph_format.space_after = Pt(6)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        b = p2.add_run(body)
        b.font.size = Pt(11)
        p2.paragraph_format.line_spacing = 1.15

    fill_cell(
        left,
        "MISSION",
        "Excellence in Every Detail",
        "To deliver world-class construction and interior fitout solutions that exceed client expectations "
        "through innovative design, precision craftsmanship, and unwavering commitment to quality and "
        "timely delivery.",
    )
    fill_cell(
        right,
        "VISION",
        "Building the Future",
        "To become the most trusted and sought-after construction and interior fitout company in the Gulf "
        "region, known for transforming spaces and setting new standards of excellence in every project we "
        "undertake.",
    )
    doc.add_paragraph()


def add_value_row(doc: Document, title: str, body: str) -> None:
    """Label + justified body (Our Values)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.space_after = Pt(6)
    t = p.add_run(title)
    t.bold = True
    t.font.size = Pt(11)
    t.font.color.rgb = ORANGE
    p.add_run("  ")
    b = p.add_run(body)
    b.font.size = Pt(11)


def add_service_table(doc: Document, services: list[tuple[str, str, str]]) -> None:
    """Numbered rows like PDF (01 … 12)."""
    tbl = doc.add_table(rows=len(services), cols=2)
    tbl.style = "Table Grid"
    for ri, (num, title, desc) in enumerate(services):
        c0, c1 = tbl.rows[ri].cells[0], tbl.rows[ri].cells[1]
        c0.width = Inches(0.55)
        c1.width = Inches(5.75)
        c0.text = ""
        c1.text = ""
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(num)
        r0.bold = True
        r0.font.size = Pt(11)
        r0.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_shading(c0, "FF8402")
        p1 = c1.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        t = p1.add_run(title)
        t.bold = True
        t.font.size = Pt(11)
        p1 = c1.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        d = p1.add_run(desc)
        d.font.size = Pt(10)
        p1.paragraph_format.space_after = Pt(4)
    doc.add_paragraph()


def add_numbered_point(doc: Document, n: int, title: str, body: str) -> None:
    """Why FARJAR — large number + title + body (PDF-style points)."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.rows[0].cells[0].width = Inches(0.5)
    tbl.rows[0].cells[1].width = Inches(5.8)
    c0, c1 = tbl.rows[0].cells[0], tbl.rows[0].cells[1]
    c0.text = ""
    c1.text = ""
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(str(n))
    r0.bold = True
    r0.font.size = Pt(16)
    r0.font.color.rgb = ORANGE
    _set_cell_shading(c0, "FFF3E8")
    p1 = c1.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    t = p1.add_run(title)
    t.bold = True
    t.font.size = Pt(11)
    p2 = c1.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    b = p2.add_run(body)
    b.font.size = Pt(10)
    p2.paragraph_format.space_after = Pt(8)


def add_table_block(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    col_widths: list[float] | None = None,
) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        _set_cell_shading(hdr[i], "E8E8E8")

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for r in p.runs:
                    r.font.size = Pt(9)

    if col_widths:
        for row in table.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)

    doc.add_paragraph()


def add_gallery_grid(doc: Document, rows: list[list[tuple[str, str]]]) -> None:
    """rows: each row is list of (image_rel_path, caption)."""
    for row in rows:
        tbl = doc.add_table(rows=1, cols=len(row))
        for ci, (rel, cap) in enumerate(row):
            cell = tbl.rows[0].cells[ci]
            cell.text = ""
            path = _root() / rel.replace("/", "\\")
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if path.is_file():
                p.add_run().add_picture(str(path), width=Inches(2.05))
            else:
                p.add_run(f"[Image: {rel}]").italic = True
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.add_run(cap)
            r2.font.size = Pt(8)
            r2.bold = True
        doc.add_paragraph()


def build() -> Path:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)

    # —— Page 1: Cover + Who we are + stats ——
    add_cover(doc)
    add_heading(doc, "Who We Are")
    add_body(
        doc,
        "FARJAR Construction & Development LLC is a trusted construction and interior fit-out company "
        "based in Oman, delivering high-quality solutions across commercial and hospitality sectors. Our "
        "expertise spans complete fit-outs, custom joinery, MEP installations, bespoke furniture, and "
        "impactful signage—each executed with precision and attention to detail.",
    )
    add_body(
        doc,
        "A key strength that sets FARJAR apart is our fully equipped in-house factory, owned and operated "
        "by us. This allows us to manage production internally, ensuring strict quality control, faster "
        "turnaround times, and cost efficiency across all projects. From custom joinery to bespoke furniture "
        "and detailed finishes, our factory enables us to maintain consistency, flexibility in design, and "
        "the ability to meet unique client requirements without relying on external suppliers.",
    )
    add_body(
        doc,
        "With a diverse portfolio that includes cafes, hotels, restaurants, museums, and corporate offices, "
        "we combine technical expertise with refined craftsmanship to bring concepts to life. Every project "
        "we undertake is driven by a commitment to accuracy, efficiency, and superior finishing.",
    )
    add_body(
        doc,
        "At FARJAR, we do not just build spaces—we engineer environments that reflect our clients' vision "
        "while meeting the highest standards of quality and performance.",
    )
    add_stat_grid(doc)

    add_page_break(doc)

    # —— Mission, Vision & Values ——
    add_heading(doc, "Mission, Vision & Values")
    add_mission_vision_block(doc)
    add_subheading(doc, "Our Values")
    add_value_row(
        doc,
        "On-Time Delivery",
        "We respect timelines and deliver projects on schedule without compromising quality. Every milestone is "
        "met with discipline and accountability.",
    )
    add_value_row(
        doc,
        "Client Centric",
        "Your vision drives our work. We listen, collaborate, and deliver spaces that exceed expectations. "
        "Client satisfaction is embedded in our culture.",
    )
    add_value_row(
        doc,
        "Quality First",
        "Every detail matters. We use premium materials and fine craftsmanship to ensure lasting excellence. "
        "No shortcuts, no compromises.",
    )

    add_page_break(doc)

    # —— Services ——
    add_heading(doc, "Our Services")
    add_body(
        doc,
        "Comprehensive solutions for every aspect of your commercial space, delivered with precision and "
        "technical excellence.",
    )
    services = [
        ("01", "Interior Design Solutions", "Creative and functional interior design tailored to commercial and residential environments."),
        ("02", "Space Planning & 3D Visualization", "Efficient space planning supported by detailed 3D concepts to help clients visualize the final outcome."),
        ("03", "Civil, Structural & Steel Works", "Execution of civil construction, structural works, and steel fabrication with a strong focus on precision and durability."),
        ("04", "MEP Services", "Complete mechanical, electrical, and plumbing solutions designed and executed to meet industry standards and project specifications."),
        ("05", "Interior Fit-Out & Finishing", "High-quality interior fit-out services including ceiling, flooring, wall finishes, and detailed finishing works."),
        ("06", "Bespoke Joinery & Carpentry", "Custom joinery and carpentry solutions crafted with precision to match design and functional requirements."),
        ("07", "Custom Furniture Solutions", "Design and manufacturing of bespoke furniture for commercial spaces, hospitality projects, and residential interiors."),
        ("08", "Retail & Display Fixtures", "Production of customized retail displays, counters, shelving systems, and brand-focused display units."),
        ("09", "F&B and Office Fit-Out", "Specialized execution of cafés, restaurants, and corporate office interiors with a focus on efficiency and high-quality finishing."),
        ("10", "Project Management & Site Supervision", "Professional project coordination, planning, and on-site supervision to ensure timely and smooth project delivery."),
        ("11", "Turnkey Interior Solutions", "End-to-end project execution from concept and design to final handover under a single point of responsibility."),
        ("12", "Renovation & Refurbishment", "Upgrade and transformation of existing spaces with minimal disruption and maximum quality."),
    ]
    add_service_table(doc, services)

    add_page_break(doc)

    # —— Projects ——
    add_heading(doc, "Projects")
    add_body(
        doc,
        "Complete portfolio of 42+ projects delivered across Oman, organized by sector.",
    )

    add_subheading(doc, "Hospitality & selected")
    add_table_block(
        doc,
        ["Project", "Location", "Scope", "Dates"],
        [
            ["4 Points by Sheraton", "Duqm Free Zone", "Fitout & Joinery", "Aug 2025 – Oct 2025"],
            ["4 Points by Sheraton", "Duqm Free Zone", "Furniture", "Oct 2025 – Feb 2026"],
            ["Damas Hotel", "Ghubra, Muscat", "Signage Installation", "2024 – 2025"],
            ["Kiriad Hotel", "Salalah", "Interior Works", "2024 – 2025"],
        ],
        [1.4, 1.35, 1.55, 1.1],
    )

    add_subheading(doc, "Food & beverage — 2023–2024")
    add_table_block(
        doc,
        ["Project", "Location", "Scope", "Dates"],
        [
            ["Fifty Five Coffee", "Airport Heights, Ghala", "Fitout, Furniture, MEP & Signage", "Aug 2023 – Nov 2023"],
            ["Fifty Five Coffee", "Mabela - 2", "AC Works", "Nov 2023"],
            ["Fifty Five Coffee", "Mawaleh, Seeb", "MEP & Civil Works", "Sep 2023 – Dec 2023"],
            ["Fifty Five Coffee", "Myriad Mall, Rusayl", "Fitout, Furniture, MEP & Signage", "Dec 2023 – Mar 2024"],
            ["Fifty Five Coffee", "Suwaiq", "AC Servicing", "Apr 2024"],
        ],
        [1.4, 1.35, 1.55, 1.1],
    )

    add_subheading(doc, "Food & beverage — 2024–2026 (selection)")
    add_table_block(
        doc,
        ["Project", "Location", "Scope", "Dates"],
        [
            ["Fifty Five Coffee", "Rubat Street, Salalah", "Fitout, Furniture, MEP & Signage", "Apr 2024 – Jun 2024"],
            ["Fifty Five Coffee", "Gardens Mall, Salalah", "Fitout, Furniture, MEP & Signage", "Aug 2024 – Oct 2024"],
            ["Fifty Five Coffee", "Haffa Walk, Salalah", "Fitout, Furniture, MEP & Signage", "Aug 2024 – Nov 2024"],
            ["Fifty Five Coffee", "Azaiba (Near Airport)", "Maintenance", "Sep 2024 – Nov 2024"],
            ["Fifty Five Coffee", "Ittin, Salalah", "Civil Works", "Oct 2024 – Dec 2024"],
            ["VIAN Croissants", "Al Khoud Square", "Furniture & Joinery", "Dec 2024 – Jan 2025"],
            ["20UR Coffee", "Sur", "Fitout, Furniture, MEP & Signage", "Dec 2024 – Jan 2025"],
            ["Jelal Efendi Restaurant", "Mawaleh, Seeb", "Fitout, Furniture, MEP & Signage", "Jan 2025 – Apr 2025"],
            ["Muscafe", "Bousher", "Fitout, MEP & Furniture", "May 2024 – Jun 2025"],
            ["Loreva (Venier Coffee)", "Grand Mall, Salalah", "Fitout, Furniture, MEP & Signage", "May 2025 – Jul 2025"],
            ["Caika Cafe", "Salalah", "Civil Works", "Jul 2025"],
            ["Italian Barrista Cafe", "Rustaq", "Fitout, Furniture, MEP & Signage", "Jul 2025 – Sep 2025"],
            ["Tea Town", "Azaiba (Near Airport)", "Furniture", "Sep 2025 – Oct 2025"],
            ["Italian Barrista Cafe", "Al Khuwair", "Fitout, Furniture, MEP & Signage", "Nov 2025 – Feb 2026"],
            ["Sign Cafe", "Al Amerat", "Fitout, Furniture, MEP & Signage", "Dec 2025 – Jan 2026"],
            ["CHAII", "Grand Flora, Salalah", "Fitout, Furniture, MEP & Signage", "Feb 2026 – Mar 2026"],
            ["CHAII", "Sahalnoot, Salalah", "Fitout, MEP & Signage", "Feb 2026 – Mar 2026"],
            ["CHAII", "Saada, Salalah", "Fitout, MEP & Signage", "Feb 2026 – Mar 2026"],
            ["Burger Gellatto", "Al Khuwair", "Fitout, Furniture, MEP & Signage", "Upcoming"],
        ],
        [1.35, 1.35, 1.55, 1.1],
    )

    add_subheading(doc, "Retail & luxury")
    add_table_block(
        doc,
        ["Project", "Location", "Scope", "Dates"],
        [
            ["Watchesco", "Qurum, Muscat", "Fitout, Furniture, MEP & Signage", "Jul 2025 – Sep 2025"],
            ["Al Rifai Roastery", "Mall of Oman, Muscat", "Fitout & Joinery", "Dec 2024 – Jan 2025"],
            ["Bariq Al Khaleej", "Wadi Al Kabir", "Kiosk", "2023 – 2024"],
            ["Bint Al Sheikh", "Muscat Exhibition Centre", "Kiosk", "Feb 2026"],
            ["Private Client", "Grand Mall, Muscat", "Joinery & Furniture", "Mar 2025 – Apr 2025"],
        ],
        [1.35, 1.35, 1.55, 1.1],
    )

    add_subheading(doc, "Culture & civic")
    add_table_block(
        doc,
        ["Project", "Location", "Scope", "Dates"],
        [
            ["Al Baleed Museum", "Salalah", "Fitout & Joinery", "Jul 2025 – Aug 2025"],
            ["National Energy Centre", "Sultanate of Oman", "Fitout, MEP & Furniture", "Jul 2024"],
            ["Bint Al Sheikh", "Salalah", "Fitout & Joinery", "Apr 2024"],
            ["Experience Oman", "Salalah", "Design", "Upcoming"],
        ],
        [1.35, 1.35, 1.55, 1.1],
    )

    add_subheading(doc, "Corporate & residential")
    add_table_block(
        doc,
        ["Project", "Location", "Scope", "Dates"],
        [
            ["Hamath Villa", "Muscat", "Fitout & Joinery", "Aug 2023 – Dec 2023"],
            ["Mancave - Bahwan", "Qurum, Muscat", "Fitout, MEP & Joinery", "Dec 2023 – Feb 2024"],
            ["Boad", "Muscat Bousher", "Custom Joinery & Signage", "Dec 2024 – Jan 2025"],
            ["Oman Oil", "Muscat", "Furniture", "May 2025"],
            ["OTE (Bahwan Group)", "Jalan Bani Bu Ali", "Furniture", "Dec 2025 – Jan 2026"],
        ],
        [1.35, 1.35, 1.55, 1.1],
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("42+ PROJECTS COMPLETED ACROSS 15+ CITIES IN OMAN")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = ORANGE
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)

    add_page_break(doc)

    # —— Gallery (images + captions like PDF) ——
    add_heading(doc, "Gallery")
    add_body(
        doc,
        "A glimpse into our craftsmanship across cafés, restaurants, museums, retail, and hospitality spaces.",
    )
    # Edit paths/captions to match your PDF; use repo JPEG/PNG assets (relative to project root).
    gallery_rows = [
        [
            ("assets/images/projects/JELAL EFENDI RESTAURANT/JELAL EFENDI RESTAURANT 01.png", "Jelal Efendi Restaurant"),
            ("assets/images/projects/AL BALEED MUSUEM/AL BALEED MUSUEM 12.jpeg", "Al Baleed Museum"),
            ("assets/images/projects/Watchesco/IMG_7100.JPG", "Watchesco"),
        ],
        [
            ("assets/images/projects/55 Azaiba/55 Azaiba 01.jpeg", "55 Coffee — Azaiba"),
            ("assets/images/projects/VIAN CROISSANTS/VIAN CROISSANTS 01.png", "VIAN Croissants"),
            ("assets/images/projects/UR 20/UR 20 (01).png", "20UR Coffee"),
        ],
        [
            ("assets/images/projects/55 Gardenes Mall/55 GARDENS MALL - SALALAH (01).png", "55 Coffee — Gardens Mall"),
            ("assets/images/projects/Muscafe/IMG_0245.png", "Muscafe"),
            ("assets/images/projects/CAIKA CAFE SALALA/CAIKA CAFE SALALA 01.png", "Caika Cafe"),
        ],
        [
            ("assets/images/projects/Myriad 55/55 MYRIAD HOTEL -RUSAYL (04).png", "55 Coffee — Myriad Mall"),
            ("assets/images/projects/Tea Town/Tea Town 01.png", "Tea Town"),
            # No Damas photo in repo; swap this path when you add assets/images/projects/Damas Hotel/...
            ("assets/images/projects/55 R_Bat/Rbat-Finalrender_page-0007.jpg", "Damas Hotel — Signage (placeholder image)"),
        ],
    ]

    add_gallery_grid(doc, gallery_rows[:3])
    add_gallery_grid(doc, [gallery_rows[3]])

    add_page_break(doc)

    # —— Why FARJAR ——
    add_heading(doc, "Why FARJAR?")
    add_body(
        doc,
        "From concept to handover, we manage every aspect of your project under one roof—design, MEP, "
        "joinery, furniture, and signage.",
    )
    why = [
        ("Proven Track Record", "42+ successfully delivered projects across cafés, hotels, restaurants, museums, retail outlets, and corporate offices."),
        ("Pan-Oman Presence", "Operating across 15+ cities including Muscat, Salalah, Sur, Suwaiq, Rustaq, and Duqm."),
        ("Quality Without Compromise", "Premium materials, skilled craftsmen, and rigorous quality control at every stage of the project."),
        ("On-Time Delivery", "Structured project management and site supervision ensure milestones are met on schedule."),
        ("End-to-End Capability", "12 specialized services covering everything from interior design and 3D visualization to renovation and refurbishment."),
        ("Single Point of Responsibility", "One coordinated team from concept to handover under one roof."),
    ]
    for i, (title, body) in enumerate(why, start=1):
        add_numbered_point(doc, i, title, body)

    add_subheading(doc, "Sectors We Serve")
    add_body(
        doc,
        "Cafés & coffee shops · Restaurants & dining · Hotels & hospitality · Retail & luxury boutiques · "
        "Museums & cultural centres · Corporate offices · Villas & residential · Kiosks & exhibition stands.",
    )

    add_page_break(doc)

    add_heading(doc, "Contact")
    add_body(
        doc,
        "Building no 164, way 3305, Dohat Al Adab street, Al Khuwair st, Muscat, Oman.",
    )
    add_body(doc, "Phone: +968 9393 0050")
    add_body(doc, "Email: contact@farjarinteriors.com")
    add_body(
        doc,
        "Instagram: @farjar.om · LinkedIn: FARJAR Construction & Development LLC.",
    )

    out_dir = _root() / "docs"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "FARJAR_Company_Profile.docx"
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    p = build()
    print(f"Wrote: {p}")
